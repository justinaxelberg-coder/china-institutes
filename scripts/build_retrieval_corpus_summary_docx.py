from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parents[1]

CORE_PATH = ROOT / "data" / "processed" / "corpus_core.csv"
PROB_PATH = ROOT / "data" / "processed" / "corpus_probable.csv"
FULL_PATH = ROOT / "outputs" / "final" / "china_institutes_corpus_full.csv"
POLICY_PATH = ROOT / "outputs" / "final" / "china_institutes_policy_corpus.csv"
PARAGRAPH_PATH = ROOT / "data" / "processed" / "corpus_paragraphs_scored.csv"

OUT_MD = ROOT / "outputs" / "final" / "retrieval_and_corpus_summary.md"
OUT_DOCX = ROOT / "outputs" / "final" / "retrieval_and_corpus_summary.docx"


def _load_frame(path: Path) -> pd.DataFrame:
    return pd.read_csv(path, low_memory=False)


def build_summary() -> dict[str, object]:
    core = _load_frame(CORE_PATH)
    probable = _load_frame(PROB_PATH)
    full = _load_frame(FULL_PATH)
    policy = _load_frame(POLICY_PATH)
    paragraphs = _load_frame(PARAGRAPH_PATH)

    return {
        "core_rows": int(len(core)),
        "probable_rows": int(len(probable)),
        "full_rows": int(len(full)),
        "policy_rows": int(len(policy)),
        "core_institutes": int(core["institute_id"].nunique()),
        "full_institutes": int(full["institute_id"].nunique()),
        "full_year_min": int(full["year"].min()),
        "full_year_max": int(full["year"].max()),
        "core_source_counts": core["source_type"].value_counts().to_dict(),
        "probable_source_counts": probable["source_type"].value_counts().to_dict(),
        "full_source_counts": full["source_type"].value_counts().to_dict(),
        "paragraph_rows": int(len(paragraphs)),
        "paragraph_docs": int(paragraphs["doc_id"].nunique()),
        "review_like_rows": int(paragraphs["is_review_like"].fillna(False).sum()),
        "dominant_domain_counts": paragraphs["dominant_domain"].value_counts().to_dict(),
        "top_core_institutes": core["institute_name"].value_counts().head(10).to_dict(),
    }


def build_markdown(summary: dict[str, object]) -> str:
    core_source = summary["core_source_counts"]
    probable_source = summary["probable_source_counts"]
    full_source = summary["full_source_counts"]
    domains = summary["dominant_domain_counts"]
    top_core = summary["top_core_institutes"]

    lines = [
        "# Current Retrieval and Corpus Summary",
        "",
        "## Scope",
        "",
        "This note describes only the retrieval and corpus layers that are currently doing the heavy lifting in the analysis. It deliberately focuses on the active production backbone and leaves out exploratory or legacy analytical layers.",
        "",
        "## What Currently Does The Heavy Lifting",
        "",
        "The present analysis is driven mainly by two retrieval streams:",
        "",
        "1. attributed **indexed publications** from the OpenAlex-based retrieval path",
        "2. a smaller attributed **grey-literature** stream from institute websites and official publication pages",
        "",
        "The main typology and within-domain orientation analysis are estimated from the **core publication corpus** built from those sources.",
        "",
        "A separate policy-document corpus is still retained, but it is not the main driver of the current publication-derived typology.",
        "",
        "## Retrieval Backbone",
        "",
        "The current production corpus is assembled from three attributed source files:",
        "",
        "- `data/interim/openalex_attributed.csv`",
        "- `data/interim/grey_attributed.csv`",
        "- `data/interim/dimensions_policy_attributed.csv`",
        "",
        "In the current workflow:",
        "",
        "- indexed publications contribute `title + abstract` text",
        "- grey literature contributes parsed page text when available, otherwise page title",
        "- policy documents are exported separately rather than folded into the main publication corpus",
        "",
        "Cross-source duplicates are removed by normalized title within institute, with preference order:",
        "",
        "1. indexed",
        "2. grey",
        "3. policy_document",
        "",
        "The `probable` corpus excludes anything already included in `core`.",
        "",
        "## Corpus Structure",
        "",
        f"- Core corpus: **{summary['core_rows']:,}** documents across **{summary['core_institutes']:,}** institutes",
        f"- Probable corpus: **{summary['probable_rows']:,}** documents",
        f"- Full publication corpus: **{summary['full_rows']:,}** documents across **{summary['full_institutes']:,}** institutes",
        f"- Separate policy corpus: **{summary['policy_rows']:,}** rows",
        f"- Year coverage in the full publication corpus: **{summary['full_year_min']}-{summary['full_year_max']}**",
        "",
        "Source composition:",
        "",
        f"- Core corpus: indexed={core_source.get('indexed', 0):,}, grey={core_source.get('grey', 0):,}",
        f"- Probable corpus: indexed={probable_source.get('indexed', 0):,}, grey={probable_source.get('grey', 0):,}",
        f"- Full publication corpus: indexed={full_source.get('indexed', 0):,}, grey={full_source.get('grey', 0):,}",
        "",
        "This means the heavy lifting is still overwhelmingly being done by the indexed publication layer, with grey literature acting as a useful but much smaller supplement.",
        "",
        "## What Actually Enters The Current Analysis",
        "",
        "The current typology/orientation rewrite is grounded in the **core publication corpus**, not the full union.",
        "",
        f"- Core documents entering the paragraph layer: **{summary['core_rows']:,}**",
        f"- Paragraph rows written for the current domain/orientation layer: **{summary['paragraph_rows']:,}**",
        f"- Documents represented in the paragraph layer: **{summary['paragraph_docs']:,}**",
        f"- Review-like paragraphs flagged: **{summary['review_like_rows']:,}**",
        "",
        "Current dominant paragraph-domain counts:",
        "",
        f"- society_culture: {domains.get('society_culture', 0):,}",
        f"- security_strategy: {domains.get('security_strategy', 0):,}",
        f"- economy_trade: {domains.get('economy_trade', 0):,}",
        f"- governance_law: {domains.get('governance_law', 0):,}",
        f"- science_technology: {domains.get('science_technology', 0):,}",
        f"- environment_climate: {domains.get('environment_climate', 0):,}",
        f"- unclassified: {domains.get('unclassified', 0):,}",
        "",
        "So, in practical terms, the current analysis is being carried mainly by the core indexed corpus, with smaller support from grey literature, and then pushed through the paragraph-level domain and orientation pipeline.",
        "",
        "## What Is Supporting Rather Than Driving",
        "",
        "- The `probable` corpus broadens descriptive coverage but is not the conservative analytical base.",
        "- The separate policy corpus is preserved for policy-facing outputs and contextual work, but it is not the backbone of the current typology.",
        "- Exploratory topic-modeling layers are not the main retrieval/corpus engine for the current analysis.",
        "",
        "## Largest Institutes In The Core Analytical Corpus",
        "",
    ]

    for institute, count in top_core.items():
        lines.append(f"- {institute}: {count:,}")

    lines.extend(
        [
            "",
            "## Bottom Line",
            "",
            "The current analytical heavy lifting is being done by a conservative core corpus built mainly from attributed OpenAlex indexed publications, supplemented by a smaller grey-literature stream, deduplicated across sources, and then carried into the paragraph-level domain/orientation pipeline. The probable corpus and policy corpus remain useful, but they are not the primary engines of the present publication-derived analysis.",
        ]
    )
    return "\n".join(lines)


def build_docx(summary: dict[str, object]) -> None:
    core_source = summary["core_source_counts"]
    probable_source = summary["probable_source_counts"]
    full_source = summary["full_source_counts"]
    domains = summary["dominant_domain_counts"]

    doc = Document()
    doc.add_heading("Current Retrieval and Corpus Summary", level=1)
    doc.add_paragraph(
        "This note describes only the retrieval and corpus layers that are currently doing the heavy lifting in the analysis. It focuses on the active production backbone and leaves out exploratory or legacy analytical layers."
    )

    doc.add_heading("What Currently Does The Heavy Lifting", level=2)
    doc.add_paragraph(
        "The present analysis is driven mainly by attributed indexed publications from the OpenAlex-based retrieval path, plus a smaller attributed grey-literature stream from institute websites and official publication pages."
    )
    doc.add_paragraph(
        "The main typology and within-domain orientation analysis are estimated from the core publication corpus built from those sources. A separate policy-document corpus is retained, but it is not the main driver of the current publication-derived typology."
    )

    doc.add_heading("Retrieval Backbone", level=2)
    retrieval_points = [
        "The current production corpus is assembled from openalex_attributed.csv, grey_attributed.csv, and dimensions_policy_attributed.csv.",
        "Indexed publications contribute title plus abstract text.",
        "Grey literature contributes parsed page text when available, otherwise page title.",
        "Policy documents are exported separately rather than folded into the main publication corpus.",
        "Cross-source duplicates are removed by normalized title within institute with preference order indexed, then grey, then policy_document.",
        "The probable corpus excludes anything already included in core.",
    ]
    for item in retrieval_points:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Corpus Structure", level=2)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    headers = ["Layer", "Rows", "Institutes", "Notes"]
    for i, header in enumerate(headers):
        t.rows[0].cells[i].text = header
    rows = [
        ("Core corpus", f"{summary['core_rows']:,}", f"{summary['core_institutes']:,}", f"indexed={core_source.get('indexed', 0):,}; grey={core_source.get('grey', 0):,}"),
        ("Probable corpus", f"{summary['probable_rows']:,}", "", f"indexed={probable_source.get('indexed', 0):,}; grey={probable_source.get('grey', 0):,}"),
        ("Full publication corpus", f"{summary['full_rows']:,}", f"{summary['full_institutes']:,}", f"indexed={full_source.get('indexed', 0):,}; grey={full_source.get('grey', 0):,}"),
        ("Separate policy corpus", f"{summary['policy_rows']:,}", "", "stored separately from the main publication corpus"),
    ]
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    doc.add_paragraph(
        f"The full publication corpus currently spans {summary['full_year_min']}-{summary['full_year_max']}. In practice, the heavy lifting is still overwhelmingly being done by the indexed publication layer, with grey literature acting as a useful but much smaller supplement."
    )

    doc.add_heading("What Actually Enters The Current Analysis", level=2)
    analysis_points = [
        f"Core documents entering the paragraph layer: {summary['core_rows']:,}",
        f"Paragraph rows written for the current domain/orientation layer: {summary['paragraph_rows']:,}",
        f"Documents represented in the paragraph layer: {summary['paragraph_docs']:,}",
        f"Review-like paragraphs flagged: {summary['review_like_rows']:,}",
    ]
    for item in analysis_points:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("Current dominant paragraph-domain counts:")
    domain_points = [
        f"society_culture: {domains.get('society_culture', 0):,}",
        f"security_strategy: {domains.get('security_strategy', 0):,}",
        f"economy_trade: {domains.get('economy_trade', 0):,}",
        f"governance_law: {domains.get('governance_law', 0):,}",
        f"science_technology: {domains.get('science_technology', 0):,}",
        f"environment_climate: {domains.get('environment_climate', 0):,}",
        f"unclassified: {domains.get('unclassified', 0):,}",
    ]
    for item in domain_points:
        doc.add_paragraph(item, style="List Bullet 2")

    doc.add_paragraph(
        "In practical terms, the current analysis is being carried mainly by the core indexed corpus, with smaller support from grey literature, and then pushed through the paragraph-level domain and orientation pipeline."
    )

    doc.add_heading("What Is Supporting Rather Than Driving", level=2)
    support_points = [
        "The probable corpus broadens descriptive coverage but is not the conservative analytical base.",
        "The separate policy corpus is preserved for policy-facing outputs and contextual work, but it is not the backbone of the current typology.",
        "Exploratory topic-modeling layers are not the main retrieval/corpus engine for the current analysis.",
    ]
    for item in support_points:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Largest Institutes In The Core Analytical Corpus", level=2)
    for institute, count in summary["top_core_institutes"].items():
        doc.add_paragraph(f"{institute}: {count:,}", style="List Bullet")

    doc.add_heading("Bottom Line", level=2)
    doc.add_paragraph(
        "The current analytical heavy lifting is being done by a conservative core corpus built mainly from attributed OpenAlex indexed publications, supplemented by a smaller grey-literature stream, deduplicated across sources, and then carried into the paragraph-level domain and orientation pipeline. The probable corpus and policy corpus remain useful, but they are not the primary engines of the present publication-derived analysis."
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)


def main() -> None:
    summary = build_summary()
    markdown = build_markdown(summary)
    OUT_MD.parent.mkdir(parents=True, exist_ok=True)
    OUT_MD.write_text(markdown, encoding="utf-8")
    build_docx(summary)
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
