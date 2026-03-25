from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parents[1]

PARAGRAPH_PATH = ROOT / "data" / "processed" / "corpus_paragraphs_scored.csv"
MODEL_DATA_PATH = ROOT / "outputs" / "tables" / "institute_typology_model_data.csv"
MAIN_MODELS_PATH = ROOT / "outputs" / "tables" / "institute_typology_main_models.csv"

OUT_MD = ROOT / "outputs" / "final" / "domain_sentiment_results_summary.md"
OUT_DOCX = ROOT / "outputs" / "final" / "domain_sentiment_results_summary.docx"


DOMAIN_LABELS = {
    "security_strategy": "Security & strategy",
    "economy_trade": "Economy & trade",
    "science_technology": "Science, technology & education",
    "environment_climate": "Environment & climate",
    "governance_law": "Governance & law",
    "society_culture": "Society & culture",
    "unclassified": "Unclassified",
}

STANCE_LABELS = {
    "neutral": "Neutral / analytic",
    "collaborator_only": "Collaborator only",
    "rival_only": "Rival only",
    "collaborator_leaning": "Collaborator leaning",
    "balanced_mixed": "Balanced mixed",
    "rival_leaning": "Rival leaning",
}


def build_summary() -> dict[str, object]:
    paragraphs = pd.read_csv(PARAGRAPH_PATH, low_memory=False)
    model_data = pd.read_csv(MODEL_DATA_PATH, low_memory=False)
    main_models = pd.read_csv(MAIN_MODELS_PATH, low_memory=False)

    top_assoc = (
        main_models.loc[:, ["outcome", "term", "estimate", "p_value"]]
        .sort_values("p_value", ascending=True)
        .head(6)
        .copy()
    )

    return {
        "paragraph_rows": int(len(paragraphs)),
        "review_like_rows": int(paragraphs["is_review_like"].fillna(False).sum()),
        "domain_counts": paragraphs["dominant_domain"].value_counts().to_dict(),
        "stance_counts": paragraphs["stance_lean"].value_counts().to_dict(),
        "institutes": int(len(model_data)),
        "low_volume": int(model_data["low_volume_flag"].fillna(False).sum()),
        "modeled_institutes": int((~model_data["low_volume_flag"].fillna(False)).sum()),
        "typology_counts": model_data["theory_typology"].value_counts().to_dict(),
        "top_assoc": top_assoc,
    }


def build_markdown(summary: dict[str, object]) -> str:
    domain_counts = summary["domain_counts"]
    stance_counts = summary["stance_counts"]
    typology_counts = summary["typology_counts"]

    lines = [
        "# Domain and Orientation Methodology With Brief Results Summary",
        "",
        "## Purpose",
        "",
        "This note summarizes the current production methodology for domain identification and orientation scoring, followed by a short summary of the current results. It reflects the active paragraph-level domain/orientation pipeline rather than older one-axis framing outputs.",
        "",
        "## 1. Domain Methodology",
        "",
        "The current production workflow identifies domains in two steps.",
        "",
        "First, the pipeline applies a rule-based China-studies thematic taxonomy. That underlying taxonomy has 15 themes, including foreign policy, trade and political economy, domestic politics, security and military affairs, Taiwan, Hong Kong, Xinjiang/Tibet, technology, environment, history and culture, Global South and regional activity, media and influence operations, society and demographics, health science, and law and institutions.",
        "",
        "Second, those 15 themes are collapsed upward into 6 broader analytical domains used in the current typology:",
        "",
        "- security_strategy",
        "- economy_trade",
        "- science_technology",
        "- environment_climate",
        "- governance_law",
        "- society_culture",
        "",
        "The key methodological shift is that the current system works at the paragraph level rather than assigning one label to an entire document. Documents from the core corpus are segmented into paragraphs, the paragraphs are scored against the thematic taxonomy, and those theme scores are then aggregated into broad-domain weights.",
        "",
        "In practice, this means a paragraph can load partly onto more than one domain rather than being forced into a single topic bucket. If no domain accumulates enough score, the paragraph is left as `unclassified`.",
        "",
        "The domain layer is therefore not a BERTopic-style latent topic model. It is a controlled, production-facing domain identification system built from the thematic taxonomy plus broad-domain cue terms and a series of cleanup rules to reduce false positives, especially around science/technology versus education, governance, and bibliographic noise.",
        "",
        "## 2. Orientation Methodology",
        "",
        "The current orientation layer is not generic positive/negative sentiment analysis. Instead, it asks a more specific question: when a paragraph discusses China in a given domain, does it frame China more as a `rival` or more as a `collaborator`?",
        "",
        "These are scored as separate dimensions rather than as opposite ends of a single line. That matters because a document or institute can be collaborative in one domain and rivalrous in another, and some passages contain both kinds of language.",
        "",
        "At the paragraph level, the orientation scorer uses:",
        "",
        "- lexicon-based matching for rival and collaborator cues",
        "- paragraph body text rather than title text",
        "- density-normalized scoring",
        "- a minimum paragraph-length threshold",
        "- suppression of review-like and metadata-thin rows",
        "",
        "The remaining domain salience not allocated to rival or collaborator is treated as `neutral`, which is best understood as neutral/analytic residual rather than as a separate positive sentiment class.",
        "",
        "The paragraph-level output is then aggregated upward to the document and institute levels, producing within-domain shares such as `security_strategy_rival_share`, `economy_trade_collaborator_share`, and `science_technology_collaborator_share`.",
        "",
        "## 3. Brief Results Summary",
        "",
        f"The current paragraph-level production layer contains **{summary['paragraph_rows']:,}** scored paragraphs. Of these, **{summary['review_like_rows']:,}** are flagged as review-like or bibliographic-style rows.",
        "",
        "### Domain distribution",
        "",
    ]

    for domain, count in domain_counts.items():
        lines.append(f"- {DOMAIN_LABELS.get(domain, domain)}: {count:,}")

    lines.extend(
        [
            "",
            "The largest broad domains at the paragraph level are currently society/culture and security/strategy, followed by economy/trade. Environment/climate is smaller but now clearly visible as its own domain rather than being buried inside political economy.",
            "",
            "### Orientation distribution",
            "",
        ]
    )

    for stance, count in stance_counts.items():
        lines.append(f"- {STANCE_LABELS.get(stance, stance)}: {count:,}")

    lines.extend(
        [
            "",
            "The most important overall pattern is that the corpus is still mostly neutral/analytic in tone, which is what we would expect from a publication corpus. Within the non-neutral share, collaborator language is more common than rival language, but mixed and rivalrous passages clearly exist and are substantial enough to support a typology richer than a single engagement-threat axis.",
            "",
            "### Institute-level typology",
            "",
            f"- Institutes in the current typology frame: **{summary['institutes']:,}**",
            f"- Low-volume institutes excluded from modeling: **{summary['low_volume']:,}**",
            f"- Institutes entering the current regression sample: **{summary['modeled_institutes']:,}**",
            "",
        ]
    )

    for typology, count in typology_counts.items():
        lines.append(f"- {typology}: {count:,}")

    lines.extend(
        [
            "",
            "Taken together, these results support the core substantive point of the rewrite: institutes do not line up neatly on a single spectrum from engagement to threat. Instead, they differ in how they distribute attention across domains and in how they orient toward China within those domains.",
            "",
            "### Brief model read",
            "",
            "The current regression layer remains secondary to the publication-derived typology and should be read associationally rather than causally. Its main value at this stage is to show that the within-domain outcomes can be modeled as distinct dependent variables rather than collapsed into one overall index.",
            "",
            "Among the current continuous models, the strongest retained associations are concentrated in security-rival and security-collaborator outcomes, which is consistent with the broader finding that security remains one of the most structuring domains in the corpus.",
        ]
    )

    return "\n".join(lines)


def build_docx(summary: dict[str, object]) -> None:
    domain_counts = summary["domain_counts"]
    stance_counts = summary["stance_counts"]
    typology_counts = summary["typology_counts"]

    doc = Document()
    doc.add_heading("Domain and Orientation Methodology With Brief Results Summary", level=1)
    doc.add_paragraph(
        "This note summarizes the current production methodology for domain identification and orientation scoring, followed by a short summary of the current results. It reflects the active paragraph-level domain/orientation pipeline rather than older one-axis framing outputs."
    )

    doc.add_heading("1. Domain Methodology", level=2)
    doc.add_paragraph(
        "The current production workflow identifies domains in two steps. First, it applies a rule-based China-studies thematic taxonomy with 15 themes. Second, those theme scores are collapsed upward into 6 broader analytical domains used in the current typology."
    )
    for item in [
        "security_strategy",
        "economy_trade",
        "science_technology",
        "environment_climate",
        "governance_law",
        "society_culture",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "The key methodological shift is that the current system works at the paragraph level rather than assigning one label to an entire document. Documents from the core corpus are segmented into paragraphs, those paragraphs are scored against the thematic taxonomy, and the resulting theme scores are aggregated into broad-domain weights."
    )
    doc.add_paragraph(
        "A paragraph can therefore load partly onto more than one domain rather than being forced into a single topic bucket. If no domain accumulates enough score, the paragraph is left as unclassified."
    )

    doc.add_heading("2. Orientation Methodology", level=2)
    doc.add_paragraph(
        "The current orientation layer is not generic positive/negative sentiment analysis. Instead, it asks whether a paragraph frames China more as a rival or more as a collaborator within the domain being discussed."
    )
    doc.add_paragraph(
        "These are scored as separate dimensions rather than as opposite ends of a single line. That design matters because a document or institute can be collaborative in one domain and rivalrous in another."
    )
    orientation_points = [
        "lexicon-based matching for rival and collaborator cues",
        "paragraph body text rather than title text",
        "density-normalized scoring",
        "a minimum paragraph-length threshold",
        "suppression of review-like and metadata-thin rows",
    ]
    for item in orientation_points:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "The remaining domain salience not allocated to rival or collaborator is treated as neutral. This neutral category should be understood as neutral/analytic residual rather than as a separate positive sentiment class."
    )
    doc.add_paragraph(
        "The paragraph-level output is then aggregated upward to the document and institute levels, producing within-domain shares such as security_strategy_rival_share, economy_trade_collaborator_share, and science_technology_collaborator_share."
    )

    doc.add_heading("3. Brief Results Summary", level=2)
    doc.add_paragraph(
        f"The current paragraph-level production layer contains {summary['paragraph_rows']:,} scored paragraphs, of which {summary['review_like_rows']:,} are flagged as review-like or bibliographic-style rows."
    )

    doc.add_heading("Domain Distribution", level=3)
    for domain, count in domain_counts.items():
        doc.add_paragraph(f"{DOMAIN_LABELS.get(domain, domain)}: {count:,}", style="List Bullet")
    doc.add_paragraph(
        "The largest broad domains at the paragraph level are currently society/culture and security/strategy, followed by economy/trade. Environment/climate is smaller but now clearly visible as its own domain."
    )

    doc.add_heading("Orientation Distribution", level=3)
    for stance, count in stance_counts.items():
        doc.add_paragraph(f"{STANCE_LABELS.get(stance, stance)}: {count:,}", style="List Bullet")
    doc.add_paragraph(
        "The most important overall pattern is that the corpus is still mostly neutral/analytic in tone, which is expected for a publication corpus. Within the non-neutral share, collaborator language is more common than rival language, but mixed and rivalrous passages are clearly present."
    )

    doc.add_heading("Institute-Level Typology", level=3)
    doc.add_paragraph(f"Institutes in the current typology frame: {summary['institutes']:,}")
    doc.add_paragraph(f"Low-volume institutes excluded from modeling: {summary['low_volume']:,}")
    doc.add_paragraph(f"Institutes entering the current regression sample: {summary['modeled_institutes']:,}")
    for typology, count in typology_counts.items():
        doc.add_paragraph(f"{typology}: {count:,}", style="List Bullet")

    doc.add_paragraph(
        "Taken together, these results support the core substantive point of the rewrite: institutes do not line up neatly on a single spectrum from engagement to threat. Instead, they differ in how they distribute attention across domains and in how they orient toward China within those domains."
    )

    doc.add_heading("Brief Model Read", level=3)
    doc.add_paragraph(
        "The current regression layer remains secondary to the publication-derived typology and should be read associationally rather than causally. Its main value at this stage is to show that the within-domain outcomes can be modeled as distinct dependent variables rather than collapsed into one overall index."
    )
    doc.add_paragraph(
        "Among the current continuous models, the strongest retained associations are concentrated in security-rival and security-collaborator outcomes, which is consistent with the broader finding that security remains one of the most structuring domains in the corpus."
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)


def main() -> None:
    summary = build_summary()
    OUT_MD.parent.mkdir(parents=True, exist_ok=True)
    OUT_MD.write_text(build_markdown(summary), encoding="utf-8")
    build_docx(summary)
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
